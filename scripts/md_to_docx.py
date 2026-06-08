"""마크다운 → Word(.docx) 변환기 (보고서 작성용).

논문 보고서를 마크다운으로 쓰고 Word 제출본으로 변환하기 위한 경량 도구.
pandoc 없이 python-docx 만으로 동작한다. 지원 문법:
  · `# / ## / ### / ####`  → Word Heading 1~4 (네비게이션 창에 목차로 잡힘)
  · `> ...`                 → 회색 인용 단락(채움 가이드용; 본문 채우면 지우면 됨)
  · `- ` / `* `             → 글머리 기호 목록
  · `1. `                   → 번호 목록
  · `---`                   → 빈 구분 단락
  · 그 외                   → 일반 단락 (인라인 **굵게** 처리)

사용:
  python scripts/md_to_docx.py docs/cp247_final_report_outline.md
  python scripts/md_to_docx.py <in.md> [out.docx]
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


def _add_inline(paragraph, text: str) -> None:
    """**굵게** 만 처리하는 최소 인라인 파서."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if i % 2 == 1:  # split 의 홀수 인덱스 = ** ** 안쪽
            run.bold = True


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.strip() == "---":
            doc.add_paragraph("")
            continue
        m = re.match(r"^(#{1,5})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # 관례: md `#`=문서 제목(Title), `##`=대분류(Heading 1), `###`=소분류(H2)…
            # python-docx 에서 level=0 이 Title 스타일.
            doc.add_heading(text, level=0 if level == 1 else level - 1)
            continue
        if line.startswith(">"):
            p = doc.add_paragraph()
            run = p.add_run(line.lstrip("> ").strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.font.size = Pt(9)
            continue
        m = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, m.group(3))
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, m.group(2))
            continue
        p = doc.add_paragraph()
        _add_inline(p, line)
    doc.save(str(docx_path))


def main() -> int:
    if len(sys.argv) < 2:
        print("usage: python scripts/md_to_docx.py <in.md> [out.docx]")
        return 2
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    md_to_docx(md_path, docx_path)
    print(f"wrote {docx_path}  ({docx_path.stat().st_size / 1024:.0f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
